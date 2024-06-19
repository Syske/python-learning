from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pandas as pd
import os
import sys
from PyQt5.QtWidgets import QApplication, QMainWindow, QPushButton, QFileDialog, QVBoxLayout, QWidget, QLabel, QTextEdit, QMessageBox
from PyQt5.QtCore import Qt

def gen_word_jumin(row):
    # 加载Word文档
    document = Document('居民模板.docx')
    # 定位到文档中的第一个表格
    table = document.tables[0]
    # 用户编号
    table.cell(0, 1).text = set_val(str(format_number(row['用户编号'])))

    # 如果需要替换多个单元格，可以添加更多的替换操作
    # 台区编号
    table.cell(1, 1).text = set_val(str(format_number(row['台区编号'])))
    # 用户名称
    table.cell(2, 1).text = set_val(format_number(row['用户名称']))
    # 用电地址
    table.cell(3, 1).text = set_val(format_number(row['用电地址']))
    # 证件号码
    table.cell(4, 1).text = set_val(str(format_number(row['证件号码'])))
    # 客户经理姓名
    table.cell(5, 1).text = set_val(format_number(row['客户经理姓名']))
    # 客户经理服务电话
    table.cell(5, 3).text = set_val(str(format_number(row['客户经理服务电话'])))

    # 联系信息
    child_table1 = table.cell(6, 0).tables[0]
    # 客户本人
    child_table1.cell(1, 1).text = set_val(format_number(row['联系人']))
    child_table1.cell(1, 2).text = '客户本人'
    child_table1.cell(1, 3).text = set_val(str(format_number(row['客户本人'])))

    # 财务联系人
    child_table1.cell(2, 1).text = set_val(format_number(row['联系人.1']))
    child_table1.cell(2, 2).text = '财务联系人'
    child_table1.cell(2, 3).text = set_val(str(format_number(row['账务联系人'])))

    # 停送电联系人
    child_table1.cell(3, 1).text = set_val(format_number(row['联系人.2']))
    child_table1.cell(3, 2).text = '停送电联系人'
    child_table1.cell(3, 3).text = set_val(str(format_number(row['停送电联系人'])))

    # 电气联系人
    child_table1.cell(4, 1).text = set_val(format_number(row['联系人.3']))
    child_table1.cell(4, 2).text = '电气联系人'
    child_table1.cell(4, 3).text = set_val(str(format_number(row['电气联系人'])))

    # 保存修改后的文档
    document.save(f'./居民/{row["用户名称"]}.docx')
    
def set_alignment_center(cell):
    for paragraph in cell.paragraphs:
        # 设置段落的对齐方式为居中
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
def format_number(num):
    # 检查数字是否为整数或小数部分为0
    return str(num).rstrip('.0')
    
def gen_word_fei_jumin(row):
    # 加载Word文档
    document = Document('非居民模板.docx')
    # 定位到文档中的第一个表格
    table = document.tables[0]
    # 用户编号
    table.cell(0, 1).text = set_val(str(format_number(row['用户编号'])))

    # 如果需要替换多个单元格，可以添加更多的替换操作
    # 台区编号
    table.cell(1, 1).text = set_val(str(format_number(row['台区编号'])))
    # 用户名称
    table.cell(2, 1).text = set_val(format_number(row['用户名称']))
    # 用电地址
    table.cell(3, 1).text = set_val(format_number(row['用电地址']))
    # 证件号码
    table.cell(4, 1).text = set_val(str(format_number(row['证件号码'])))
    # 客户经理姓名
    table.cell(5, 1).text = set_val(format_number(row['客户经理姓名']))
    # 客户经理服务电话
    table.cell(5, 3).text = set_val(str(format_number(row['客户经理服务电话'])))
    
    # 证件信息
    child_table = table.cell(6, 0).tables[0]
    
    if format_number(row['营业执照']) != None and format_number(row['营业执照']) != 'nan':
        child_table.cell(1, 1).text = '☑'
        child_table.cell(1, 3).text = str(format_number(row['营业执照']))
        # 设置段落的对齐方式为居中
        set_alignment_center(child_table.cell(1, 1))
        
    if format_number(row['事业单位法人证书']) != None and format_number(row['事业单位法人证书']) != 'nan':
        child_table.cell(2, 1).text = '☑'
        child_table.cell(2, 3).text = str(format_number(row['事业单位法人证书']))
        # 设置段落的对齐方式为居中
        set_alignment_center(child_table.cell(2, 1))
    
    if format_number(row['统一社会信用代码证']) != None and format_number(row['统一社会信用代码证']) != 'nan':
        child_table.cell(3, 1).text = '☑'
        child_table.cell(3, 3).text = str(format_number(row['统一社会信用代码证']))
        # 设置段落的对齐方式为居中
        set_alignment_center(child_table.cell(3, 1))
        
    if format_number(row['税务登记证']) != None and format_number(row['税务登记证']) != 'nan':
        child_table.cell(4, 1).text = '☑'
        child_table.cell(4, 3).text = str(format_number(row['税务登记证']))
        # 设置段落的对齐方式为居中
        set_alignment_center(child_table.cell(4, 1))
        
    if format_number(row['身份证']) != 'nan':
        child_table.cell(5, 3).text = str(format_number(row['身份证']))

    child_table1 = table.cell(6, 0).tables[1]
    # 联系信息
    # 客户本人
    child_table1.cell(1, 1).text = set_val(format_number(row['联系人']))
    child_table1.cell(1, 2).text = '客户本人'
    child_table1.cell(1, 3).text = set_val(str(format_number(row['客户本人'])))

    # 财务联系人
    child_table1.cell(2, 1).text = set_val(format_number(row['联系人.1']))
    child_table1.cell(2, 2).text = '财务联系人'
    child_table1.cell(2, 3).text = set_val(str(format_number(row['账务联系人'])))

    # 停送电联系人
    child_table1.cell(3, 1).text = set_val(format_number(row['联系人.2']))
    child_table1.cell(3, 2).text = '停送电联系人'
    child_table1.cell(3, 3).text = set_val(str(format_number(row['停送电联系人'])))

    # 电气联系人
    child_table1.cell(4, 1).text = set_val(format_number(row['联系人.3']))
    child_table1.cell(4, 2).text = '电气联系人'
    child_table1.cell(4, 3).text = set_val(str(format_number(row['电气联系人'])))

    # 保存修改后的文档
    document.save(f'./非居民/{row["用户名称"]}.docx')
    
def set_val(val):
    if val == 'nan' or val == None:
        return ''
    return val
    
def gen_word(row):
    if not os.path.exists('./非居民'):
        os.mkdir('./非居民')
        
    if not os.path.exists('./居民'):
        os.mkdir('./居民')
        
    if (format_number(row['营业执照']) == None or format_number(row['营业执照']) == 'nan') and (format_number(row['事业单位法人证书']) == None or format_number(row['事业单位法人证书']) == 'nan') and (format_number(row['统一社会信用代码证']) == None or format_number(row['统一社会信用代码证']) == 'nan') and (format_number(row['税务登记证']) == None or format_number(row['税务登记证']) == 'nan') and (format_number(row['身份证']) == None or format_number(row['身份证']) == 'nan'):
        gen_word_jumin(row)
    else:
        gen_word_fei_jumin(row)
    
def read_excel(fileName):
    # 读取名为'Sheet1'的工作表
    df = pd.read_excel(f'{fileName}', sheet_name='Sheet1')
    df = df.rename(columns={
    '非居民客户': '营业执照',
    'Unnamed: 8': '事业单位法人证书',
    'Unnamed: 9': '统一社会信用代码证',
    'Unnamed: 10': '税务登记证',
    'Unnamed: 11': '身份证'
})
    df = df.iloc[2:]
    for index, row in df.iterrows():
        print(type(row)) 
        gen_word(row)

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.initUI()

    def initUI(self):
        # 创建按钮来选择文件并解析内容
        self.choose_button = QPushButton('选择并解析文件', self)
        self.choose_button.clicked.connect(self.chooseAndParseFile)

        # 创建文本编辑框来显示文件内容
        self.text_edit = QTextEdit(self)
        self.text_edit.setReadOnly(True)

        # 设置布局
        layout = QVBoxLayout()
        layout.addWidget(self.choose_button)
        layout.addWidget(self.text_edit)

        central_widget = QWidget()
        central_widget.setLayout(layout)
        self.setCentralWidget(central_widget)

        self.setGeometry(300, 300, 500, 400)
        self.setWindowTitle('文件解析器')
        self.show()

    def chooseAndParseFile(self):
        # 打开文件选择对话框
        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        fileName, _ = QFileDialog.getOpenFileName(self, "选择文件", "", "Excel Files (*.xlsx)", options=options)

        # 如果用户选择了文件，则解析并显示内容
        if fileName:
            try:
                # 使用pandas读取文件内容（这里假设是CSV格式）
                print(fileName)
                read_excel(fileName)
                QMessageBox.information(self, "提示", "执行完成！")
            except Exception as e:
                # 如果解析出错，显示错误消息
                self.text_edit.setText(str(e))
                
if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = MainWindow()
    sys.exit(app.exec_())
    # read_excel("客户信息.xlsx")